#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
小财进项目文档 - Markdown 转 Word（简化版）
"""
import sys
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def convert_md_to_word():
    """转换 Markdown 到 Word"""

    # 文档路径
    md_file = "docs/PROJECT_DOCUMENTATION.md"
    output_file = "docs/小财进项目文档.docx"

    print("=" * 70)
    print("📄 小财进项目文档 - Markdown 转 Word")
    print("=" * 70)
    print(f"📖 输入文件: {md_file}")
    print(f"📝 输出文件: {output_file}")
    print()

    # 读取 Markdown 文件
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"❌ 错误：找不到文件 {md_file}")
        return

    # 创建 Word 文档
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    # 标题样式
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '微软雅黑'
        heading_font.bold = True
        heading_font.size = Pt(20 - i * 2)
        heading_font.color.rgb = RGBColor(51, 51, 51)

    # 解析内容
    in_code_block = False
    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        line = line.rstrip()

        # 跳过空行
        if not line:
            if not in_code_block and not in_table:
                doc.add_paragraph()
            continue

        # 代码块
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            para = doc.add_paragraph(line)
            para.runs[0].font.name = 'Consolas'
            para.runs[0].font.size = Pt(9)
            continue

        # 表格检测
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
            else:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells and cells[0]:
                    table_rows.append(cells)
            continue

        if in_table:
            if not line.startswith('|'):
                # 表格结束
                if table_headers and table_rows:
                    # 添加表格
                    table = doc.add_table(rows=1 + len(table_rows), cols=len(table_headers))
                    table.style = 'Light Grid Accent 1'

                    # 表头
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(table_headers):
                        cell = header_cells[i]
                        cell.text = header
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # 表格内容
                    for i, row in enumerate(table_rows):
                        row_cells = table.rows[i + 1].cells
                        for j, cell_text in enumerate(row):
                            cell = row_cells[j]
                            cell.text = str(cell_text)

                table_headers = []
                table_rows = []
                in_table = False
            continue

        # 标题
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=level)
            continue

        # 列表
        if line.strip().startswith('- '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')
            continue

        # 普通段落
        # 处理粗体
        if '**' in line:
            parts = line.split('**')
            para = doc.add_paragraph()
            for i, part in enumerate(parts):
                run = para.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        else:
            doc.add_paragraph(line)

    # 保存文档
    try:
        doc.save(output_file)
        print("✅ 转换成功！")
        print(f"\n📁 文件位置: {output_file}")
        print("\n💡 提示：你可以直接双击打开这个 .docx 文件")
    except Exception as e:
        print(f"\n❌ 保存失败: {e}")

if __name__ == "__main__":
    convert_md_to_word()
